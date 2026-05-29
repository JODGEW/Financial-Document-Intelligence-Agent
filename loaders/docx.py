"""DOCX format handler.

Uses python-docx directly (rather than ``Docx2txtLoader``) so we can preserve
heading hierarchy and emit table contents in the same ``[Extracted Tables]``
format the PDF handler uses.
"""

from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document

from .registry import FormatHandler, register
from .tables import extract_docx_tables


_HEADING_STYLE_PREFIX = "heading"


def _heading_level(style_name: str) -> int | None:
    """Return 1/2/3 for Word heading styles, else None."""
    name = (style_name or "").lower().strip()
    if not name.startswith(_HEADING_STYLE_PREFIX):
        return None
    suffix = name[len(_HEADING_STYLE_PREFIX) :].strip()
    if not suffix:
        return None
    try:
        level = int(suffix.split()[0])
    except ValueError:
        return None
    return level if 1 <= level <= 3 else None


def _section_title(stack: list[tuple[int, str]]) -> str | None:
    if not stack:
        return None
    return " > ".join(text for _level, text in stack)


def _load(path: str) -> list[Document]:
    from docx import Document as DocxDocument  # python-docx

    docx_doc = DocxDocument(path)
    heading_stack: list[tuple[int, str]] = []
    sections: list[tuple[str | None, list[str]]] = [(None, [])]

    for paragraph in docx_doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        level = _heading_level(getattr(paragraph.style, "name", "") or "")
        if level is not None:
            while heading_stack and heading_stack[-1][0] >= level:
                heading_stack.pop()
            heading_stack.append((level, text))
            sections.append((_section_title(heading_stack), []))
            continue
        sections[-1][1].append(text)

    docs: list[Document] = []
    for title, paragraphs in sections:
        body = "\n".join(paragraphs).strip()
        if not body:
            continue
        metadata: dict = {"source": path}
        if title:
            metadata["section_title"] = title
        docs.append(Document(page_content=body, metadata=metadata))

    if not docs:
        # Fall back to one Document with the joined plain text so empty-heading
        # documents still flow through the pipeline.
        full_text = "\n".join(p.text for p in docx_doc.paragraphs).strip()
        if full_text:
            docs.append(Document(page_content=full_text, metadata={"source": path}))

    table_blocks = extract_docx_tables(path)
    if table_blocks and docs:
        joined = "\n\n".join(table_blocks)
        last = docs[-1]
        last.page_content = (
            f"{last.page_content.rstrip()}\n\n[Extracted Tables]\n{joined}"
        )
        last.metadata["contains_tables"] = True

    return docs


def _split(docs: list[Document]) -> list[Document]:
    from ingest import _recursive_splitter

    splitter = _recursive_splitter(chunk_size=900, chunk_overlap=100)
    chunks: list[Document] = []
    for doc in docs:
        for piece in splitter.split_documents([doc]):
            piece.metadata = {**doc.metadata, **piece.metadata}
            chunks.append(piece)
    return chunks


register(
    FormatHandler(
        extensions=(".docx",),
        loader=_load,
        splitter=_split,
        format_family="text",
        extract_tables=True,
    )
)
